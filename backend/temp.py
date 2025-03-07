from docx import Document

# Create a new Word document
doc = Document()

# Add heading (optional)
doc.add_paragraph("Subhash Bangalore Satheesha\n"
                  "1170 Tripoli St, Riverside, CA 92507\n"
                  "sbang030@ucr.edu | +1 (951) 377-3937\n"
                  "[LinkedIn](https://www.linkedin.com/in/subhash-b-s/)\n")

doc.add_paragraph("\nFeb 21, 2025\n")

doc.add_paragraph("Psychology Consultation Specialists, PLLC\n"
                  "Riverside, CA\n")

doc.add_paragraph("Subject: Application for Early Morning Walking Companion Position\n")

# Add the body of the letter
doc.add_paragraph("Dear Hiring Team,\n")

doc.add_paragraph("I am excited to apply for the UCR Campus Early Morning Walking Companion role. "
                  "As a graduate student at the University of California, Riverside, I appreciate the importance of consistency, "
                  "companionship, and community support—values that align closely with this opportunity.\n")

doc.add_paragraph("I am a responsible, punctual, and sociable individual who enjoys engaging conversations and outdoor activities. "
                  "My background as a former sports captain and national-level hockey player has ingrained in me the discipline "
                  "and reliability needed for this role. Additionally, my experience leading student initiatives has strengthened "
                  "my ability to foster positive and engaging interactions.\n")

doc.add_paragraph("Living near UCR, I am well-positioned to commit to early morning walks and ensure a dependable presence. "
                  "I would be delighted to contribute to a friendly and uplifting walking experience.\n")

doc.add_paragraph("Thank you for your time and consideration. I look forward to the possibility of assisting in this meaningful role.\n")

doc.add_paragraph("Best regards,\n")
doc.add_paragraph("Subhash Bangalore Satheesha")

# Save the document
file_path = "/mnt/data/UCR_Walking_Companion_Cover_Letter.docx"
doc.save(file_path)

file_path
